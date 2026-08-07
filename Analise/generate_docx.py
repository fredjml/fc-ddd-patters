from docx import Document
import sys

md_path = 'Analise/relatorio-executivo.md'
docx_path = 'Analise/relatorio-executivo.docx'

def add_paragraph_from_line(doc, line):
    if line.startswith('# '):
        doc.add_heading(line[2:].strip(), level=1)
    elif line.startswith('## '):
        doc.add_heading(line[3:].strip(), level=2)
    elif line.startswith('### '):
        doc.add_heading(line[4:].strip(), level=3)
    elif line.startswith('- '):
        doc.add_paragraph(line[2:].strip(), style='List Bullet')
    elif line.strip() == '---':
        doc.add_page_break()
    else:
        doc.add_paragraph(line)


def main():
    doc = Document()
    with open(md_path, 'r', encoding='utf-8') as f:
        for raw in f:
            add_paragraph_from_line(doc, raw.rstrip('\n'))
    doc.save(docx_path)
    print('Saved', docx_path)

if __name__ == '__main__':
    main()
